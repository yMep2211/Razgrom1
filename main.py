import sys
from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTableWidget, QTableWidgetItem, QComboBox,
    QDateEdit, QLabel, QFrame, QLineEdit, QMessageBox, QDialog, QSpinBox, QFormLayout, QHeaderView, QTextEdit, QFileDialog,
)
from PyQt6.QtCore import QDate
from PyQt6.QtGui import QDoubleValidator
from db import SessionLocal
from models import Пользователь, ПлатежиПользователей, Платеж, Категория
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen.canvas import Canvas
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily
from reportlab.lib.enums import TA_CENTER
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class LoginWindow(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Авторизация")
        self.setFixedSize(400, 200)
        
        self.db = SessionLocal()  
        
        layout = QVBoxLayout(self)
        form_layout = QFormLayout()
        
        self.login_combo = QComboBox()
        self.load_users()
        
        self.password_edit = QLineEdit()
        self.password_edit.setEchoMode(QLineEdit.EchoMode.Password)
                
        buttons_layout = QHBoxLayout()
        self.btn_login = QPushButton("Войти")
        self.btn_exit = QPushButton("Выход")
        buttons_layout.addWidget(self.btn_login)
        buttons_layout.addWidget(self.btn_exit)
        
        form_layout.addRow("Имя пользователя", self.login_combo)
        form_layout.addRow("Пароль", self.password_edit)
        
        layout.addLayout(form_layout)
        layout.addLayout(buttons_layout)
        
        self.btn_login.clicked.connect(self.authenticate)
        self.btn_exit.clicked.connect(self.reject)
    
    def load_users(self):
        try:
            users = self.db.query(Пользователь).all()
            for user in users:
                self.login_combo.addItem(user.логин, user)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить пользователей: {str(e)}")
    
    def authenticate(self):
        try:
            user = self.login_combo.currentData()
            if not user:
                QMessageBox.warning(self, "Ошибка", "Выберите пользователя")
                return
                
            if user.пароль == self.password_edit.text():
                self.current_user = user
                self.accept()
            else:
                QMessageBox.warning(self, "Ошибка", "Неверный пароль")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка авторизации: {str(e)}")
    
    def closeEvent(self, event):
        self.db.close()
        super().closeEvent(event)

class AddPayment(QDialog):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Добавление платежа")
        self.setFixedSize(400, 250)
        
        self.db = SessionLocal()
        
        form_layout = QFormLayout()
        
        self.category_combo = QComboBox()
        self.load_categories()
        form_layout.addRow("Категория:", self.category_combo)
        
        self.payment_name = QLineEdit()
        self.payment_name.textChanged.connect(self.validate_inputs)
        form_layout.addRow("Назначение платежа:", self.payment_name)
        
        self.quantity = QSpinBox()
        self.quantity.setMinimum(1)
        self.quantity.setValue(1)
        form_layout.addRow("Количество:", self.quantity)
        
        self.price = QLineEdit()
        self.price.setValidator(QDoubleValidator(0.01, 999999.99, 2))
        self.price.textChanged.connect(self.validate_inputs)
        price_layout = QHBoxLayout()
        price_layout.addWidget(self.price)
        price_layout.addWidget(QLabel("р."))
        price_layout.addStretch()
        form_layout.addRow("Цена:", price_layout)
        
        buttons_layout = QHBoxLayout()
        self.btn_add = QPushButton("Добавить")  
        self.btn_cancel = QPushButton("Отменить")
        buttons_layout.addWidget(self.btn_add)
        buttons_layout.addWidget(self.btn_cancel)
        
        form_layout.addRow(buttons_layout)
        self.setLayout(form_layout)
        
        self.btn_add.clicked.connect(self.try_accept)
        self.btn_cancel.clicked.connect(self.reject)
    
    def validate_inputs(self):
        name_valid = self.is_valid_payment_name(self.payment_name.text())
        price_valid = self.is_valid_price(self.price.text())
        self.btn_add.setEnabled(name_valid and price_valid)
    
    def is_valid_payment_name(self, text: str) -> bool:
        
        clean_text = text.replace(" ", "")
        if len(clean_text) < 3:
            return False
            
        russian_letters = sum(1 for c in clean_text.lower() if 'а' <= c <= 'я')
        has_english = any(c.isalpha() and c.isascii() for c in clean_text)
        
        return russian_letters >= 3 and not has_english
    
    def is_valid_price(self, text: str) -> bool:
        try:
            return float(text.replace(",", ".")) > 0
        except ValueError:
            return False
    
    def try_accept(self):
        if not self.is_valid_payment_name(self.payment_name.text()):
            QMessageBox.warning(self, "Ошибка", 
                "Назначение платежа должно:\n"
                "- Содержать минимум 3 русские буквы\n")
            return
            
        if not self.is_valid_price(self.price.text()):
            QMessageBox.warning(self, "Ошибка", 
                "Цена должна быть положительным числом")
            return
            
        self.accept()
    
    def get_payment_data(self):
        return {
            'name': self.payment_name.text().strip(),
            'quantity': self.quantity.value(),
            'price': float(self.price.text().replace(",", ".")),
            'category_id': self.category_combo.currentData()
        }
    
    def load_categories(self):
        try:
            categories = self.db.query(Категория).order_by(Категория.название).all()
            for category in categories:
                self.category_combo.addItem(category.название, category.id)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить категории: {str(e)}")
    
    def closeEvent(self, event):
        self.db.close()
        super().closeEvent(event)
        
class DeletePayment(QDialog):
    def __init__(self, current_user, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Удаление платежа")
        self.setFixedSize(500, 300)
        self.current_user = current_user
        self.db = SessionLocal()
        
        layout = QVBoxLayout()
        
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels([
            "ID", 
            "Наименование", 
            "Дата", 
            "Сумма", 
            "Категория"
        ])
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setSelectionMode(QTableWidget.SelectionMode.SingleSelection)
        
        self.load_payments()
        layout.addWidget(self.table)
        
        buttons_layout = QHBoxLayout()
        self.btn_delete = QPushButton("Удалить выбранное")
        self.btn_cancel = QPushButton("Отмена")
        buttons_layout.addWidget(self.btn_delete)
        buttons_layout.addWidget(self.btn_cancel)
        
        layout.addLayout(buttons_layout)
        self.setLayout(layout)
        
        self.btn_delete.clicked.connect(self.confirm_delete)
        self.btn_cancel.clicked.connect(self.reject)
    
    def load_payments(self):
        try:
            payments = self.db.query(Платеж, Категория)\
                .join(ПлатежиПользователей, ПлатежиПользователей.id_платежа == Платеж.id)\
                .join(Категория, Категория.id == Платеж.id_категории)\
                .filter(ПлатежиПользователей.id_пользователя == self.current_user.id)\
                .order_by(Платеж.дата.desc())\
                .all()
            
            self.table.setRowCount(len(payments))
            
            for row, (payment, category) in enumerate(payments):
                self.table.setItem(row, 0, QTableWidgetItem(str(payment.id)))
                self.table.setItem(row, 1, QTableWidgetItem(payment.наименование_платежа))
                self.table.setItem(row, 2, QTableWidgetItem(payment.дата.strftime("%d.%m.%Y")))
                self.table.setItem(row, 3, QTableWidgetItem(f"{payment.стоимость:.2f}"))
                self.table.setItem(row, 4, QTableWidgetItem(category.название))
            
            self.table.resizeColumnsToContents()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить платежи: {str(e)}")
    
    def confirm_delete(self):
        selected_row = self.table.currentRow()
        if selected_row == -1:
            QMessageBox.warning(self, "Ошибка", "Выберите платеж для удаления")
            return

        payment_id = int(self.table.item(selected_row, 0).text())
        payment_name = self.table.item(selected_row, 1).text()

        reply = QMessageBox.question(
            self,
            "Подтверждение удаления",
            f"Вы действительно хотите удалить платеж?\nID: {payment_id}\nНаименование: {payment_name}",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
        )

        if reply == QMessageBox.StandardButton.Yes:
            try:

                self.db.query(ПлатежиПользователей)\
                    .filter(ПлатежиПользователей.id_платежа == payment_id)\
                    .delete()
                    
                self.db.query(Платеж)\
                    .filter(Платеж.id == payment_id)\
                    .delete()
                    
                self.db.commit()
                self.load_payments()  

            except Exception as e:
                self.db.rollback()
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить платеж: {str(e)}")

    
    def closeEvent(self, event):
        self.db.close()
        super().closeEvent(event)

class PaymentApp(QMainWindow):
    def __init__(self, current_user):
        super().__init__()
        self.current_user = current_user  
        self.setWindowTitle(f"Управление платежами - {current_user.логин} ({current_user.фамилия} {current_user.имя})")
        self.setGeometry(100, 100, 800, 600)
        
        
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        
        layout = QVBoxLayout()
        main_widget.setLayout(layout)
        
        control_panel = QHBoxLayout()
        
        self.btn_add = QPushButton("+")
        self.btn_add.clicked.connect(self.add_payment)
        self.btn_remove = QPushButton("-")
        self.btn_remove.clicked.connect(self.delete_payment)

        #Разделитель (Вот эта палочка на макете)
        divider = QFrame()
        divider.setFrameShape(QFrame.Shape.VLine)
        
        date_layout = QHBoxLayout()
        
        self.date_from = QDateEdit(QDate.currentDate())
        self.date_from.setCalendarPopup(True)
        self.date_from.setDisplayFormat("dd.MM.yyyy")
        self.date_from.setDate(QDate(2016, 11, 1))
        
        self.date_to = QDateEdit(QDate.currentDate())
        self.date_to.setCalendarPopup(True)
        self.date_to.setDisplayFormat("dd.MM.yyyy")
        
        date_layout.addWidget(QLabel("с"))
        date_layout.addWidget(self.date_from)
        date_layout.addWidget(QLabel("по"))
        date_layout.addWidget(self.date_to)
        
        category_layout = QHBoxLayout()
        category_layout.addWidget(QLabel("Категория:"))
        
        self.category_combo = QComboBox()
        self.load_categories_to_combobox()
        
        self.btn_select = QPushButton("Выбрать")
        self.btn_select.clicked.connect(self.apply_filters)
        self.btn_clear = QPushButton("Очистить")
        self.btn_clear.clicked.connect(self.clear_filters)
        
        self.btn_report = QPushButton("Отчет")
        self.btn_report.clicked.connect(self.generate_report)

        
        control_panel.addWidget(self.btn_add)
        control_panel.addWidget(self.btn_remove)
        control_panel.addWidget(divider)
        control_panel.addLayout(date_layout)
        control_panel.addSpacing(20)
        control_panel.addLayout(category_layout)
        category_layout.addWidget(self.category_combo)
        category_layout.addWidget(self.btn_select)
        category_layout.addWidget(self.btn_clear)
        category_layout.addWidget(self.btn_report)
        
        self.table = QTableWidget()
        self.table.setColumnCount(5)
        self.table.setHorizontalHeaderLabels([
            "Наименование платежа",
            "Количество", 
            "Цена",
            "Сумма",
            "Категория"
        ])
        
        self.setup_table_size()
        
        self.load_payments()
        
        layout.addLayout(control_panel)
        layout.addWidget(self.table)
        
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
    
    def setup_table_size(self):
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)  
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.ResizeToContents)  
        header.setSectionResizeMode(2, QHeaderView.ResizeMode.ResizeToContents)  
        header.setSectionResizeMode(3, QHeaderView.ResizeMode.ResizeToContents)  
        header.setSectionResizeMode(4, QHeaderView.ResizeMode.ResizeToContents)  
    
    def load_payments(self):
        self.clear_filters()
        try:
            db = SessionLocal()
            
            payments = db.query(Платеж, Категория)\
                .join(ПлатежиПользователей, ПлатежиПользователей.id_платежа == Платеж.id)\
                .join(Категория, Категория.id == Платеж.id_категории)\
                .filter(ПлатежиПользователей.id_пользователя == self.current_user.id)\
                .order_by(Платеж.дата.desc())\
                .all()
            
            self.table.setRowCount(len(payments))
            
            for row, (payment, category) in enumerate(payments):
                self.table.setItem(row, 0, QTableWidgetItem(payment.наименование_платежа))
                self.table.setItem(row, 1, QTableWidgetItem(f"{payment.количество:.2f}"))
                self.table.setItem(row, 2, QTableWidgetItem(f"{payment.цена:.2f}"))
                self.table.setItem(row, 3, QTableWidgetItem(f"{payment.стоимость:.2f}"))
                self.table.setItem(row, 4, QTableWidgetItem(category.название))
                
            self.table.resizeColumnsToContents()
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить платежи: {str(e)}")
        finally:
            db.close()

    def add_payment(self):
        dialog = AddPayment()
        if dialog.exec() == QDialog.DialogCode.Accepted:
            try:
                db = SessionLocal()
                data = dialog.get_payment_data()
                
                new_payment = Платеж(
                    дата=QDate.currentDate().toPyDate(),
                    id_категории=data['category_id'],
                    наименование_платежа=data['name'],
                    количество=data['quantity'],
                    цена=data['price']
                )
                
                db.add(new_payment)
                db.flush()
                
                db.add(ПлатежиПользователей(
                    id_пользователя=self.current_user.id,
                    id_платежа=new_payment.id
                ))
                db.commit()
                
                QMessageBox.information(self, "Успех", "Платеж добавлен!")
                self.load_payments()
                
            except Exception as e:
                db.rollback()
                QMessageBox.critical(self, "Ошибка", f"Ошибка добавления: {str(e)}")
            finally:
                db.close()
                
    def delete_payment(self):
        dialog = DeletePayment(self.current_user, self)
        dialog.exec()  
        self.load_payments()  

    def load_categories_to_combobox(self):
        try:
            db = SessionLocal()
            categories = db.query(Категория).order_by(Категория.название).all()
            
            self.category_combo.clear()
            self.category_combo.addItem("Все категории", None)  
            
            for category in categories:
                self.category_combo.addItem(category.название, category.id)
                
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось загрузить категории: {str(e)}")
        finally:
            db.close()
            
    def apply_filters(self):
        try:
            if self.date_from.date() > self.date_to.date():
                QMessageBox.warning(self, "Ошибка", "Дата 'с' не может быть больше даты 'по'")
                return
                
            db = SessionLocal()
            
            date_from = self.date_from.date().toPyDate()
            date_to = self.date_to.date().toPyDate()
            category_id = self.category_combo.currentData()
            
            query = db.query(Платеж, Категория)\
                .join(ПлатежиПользователей, ПлатежиПользователей.id_платежа == Платеж.id)\
                .join(Категория, Категория.id == Платеж.id_категории)\
                .filter(ПлатежиПользователей.id_пользователя == self.current_user.id)\
                .filter(Платеж.дата.between(date_from, date_to))
                
            if category_id:
                query = query.filter(Платеж.id_категории == category_id)
                
            payments = query.order_by(Платеж.дата.desc()).all()
            
            self.update_table(payments)
            self.setup_table_size()
            
            
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось применить фильтры: {str(e)}")
        finally:
            db.close()
                
    def clear_filters(self):
        today = QDate.currentDate()
        self.date_from.setDate(QDate(2016, 11, 1))
        self.date_to.setDate(today)
        
        self.category_combo.setCurrentIndex(0)
        
        self.apply_filters()
        self.setup_table_size()
        
    def update_table(self, payments):
        self.table.setRowCount(len(payments))
        
        for row, (payment, category) in enumerate(payments):
            self.table.setItem(row, 0, QTableWidgetItem(payment.наименование_платежа))
            self.table.setItem(row, 1, QTableWidgetItem(f"{payment.количество:.2f}"))
            self.table.setItem(row, 2, QTableWidgetItem(f"{payment.цена:.2f}"))
            self.table.setItem(row, 3, QTableWidgetItem(f"{payment.стоимость:.2f}"))
            self.table.setItem(row, 4, QTableWidgetItem(category.название))
            
    def generate_report(self):
        try:
            db = SessionLocal()

            report_data = self.prepare_report_data(db)
            if not report_data:
                return

            self.show_report_preview(report_data)

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка генерации отчета:\n{str(e)}")
        finally:
            db.close()
            
    def prepare_report_data(self, db):
        date_from = self.date_from.date().toPyDate()
        date_to = self.date_to.date().toPyDate()
        category_id = self.category_combo.currentData()

        query = db.query(Платеж, Категория)\
            .join(ПлатежиПользователей, ПлатежиПользователей.id_платежа == Платеж.id)\
            .join(Категория, Категория.id == Платеж.id_категории)\
            .filter(ПлатежиПользователей.id_пользователя == self.current_user.id)\
            .filter(Платеж.дата.between(date_from, date_to))

        if category_id:
            query = query.filter(Платеж.id_категории == category_id)

        payments = query.order_by(Категория.название, Платеж.дата).all()

        if not payments:
            QMessageBox.information(self, "Информация", "Нет данных для отчета в выбранном периоде")
            return None

        grouped = {}
        total_sum = 0

        for payment, category in payments:
            grouped.setdefault(category.название, []).append(payment)
            total_sum += float(payment.стоимость)

        return {
            'date_from': date_from,
            'date_to': date_to,
            'grouped_payments': grouped,
            'total_sum': total_sum,
            'user': self.current_user
        }
        
    def show_report_preview(self, report_data):
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle("Предпросмотр отчета")
        preview_dialog.resize(800, 600)

        layout = QVBoxLayout(preview_dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        
        report_text = self.generate_report_text(report_data)
        text_edit.setHtml(report_text)
        
        layout.addWidget(text_edit)

        buttons_layout = QHBoxLayout()
        
        btn_pdf = QPushButton("Сохранить в PDF")
        btn_pdf.clicked.connect(lambda: self.save_to_pdf(report_data))
        
        btn_docx = QPushButton("Сохранить в DOCX")
        btn_docx.clicked.connect(lambda: self.save_to_docx(report_data))
        
        btn_close = QPushButton("Закрыть")
        btn_close.clicked.connect(preview_dialog.reject)
        
        buttons_layout.addWidget(btn_pdf)
        buttons_layout.addWidget(btn_docx)
        buttons_layout.addWidget(btn_close)
        
        layout.addLayout(buttons_layout)

        preview_dialog.exec()
        
    def generate_report_text(self, report_data):
        html = f"""
        <html>
        <head>
        <style>
        body {{ font-family: Arial; margin: 20px; }}
        h1 {{ text-align: center; }}
        h2 {{ margin-top: 20px; }}
        .payment {{ margin-left: 20px; margin-bottom: 5px; }}
        .total {{ font-weight: bold; margin-top: 20px; }}
        .footer {{ margin-top: 40px; font-size: smaller; }}
        </style>
        </head>
        <body>
        <h1>Отчет о платежах</h1>
        <p style="text-align: center;">
            Период: {report_data['date_from'].strftime('%d.%m.%Y')} - {report_data['date_to'].strftime('%d.%m.%Y')}
        </p>
        """

        for category, payments in report_data['grouped_payments'].items():
            html += f"<h2>{category}</h2>"
            for payment in payments:
                html += f"""
                <div class="payment">
                    {payment.дата.strftime('%d.%m.%Y')} - {payment.наименование_платежа} 
                    <span style="float: right;">{payment.стоимость:.2f} р.</span>
                </div>
                """
        
        html += f"""
        <div class="total">ИТОГО: {report_data['total_sum']:.2f} р.</div>
        <div class="footer">
            {report_data['user'].фамилия} {report_data['user'].имя} {report_data['user'].отчество or ''}
        </div>
        </body>
        </html>
        """
        
        return html
    
    def save_to_pdf(self, report_data):
        try:
            filename, _ = QFileDialog.getSaveFileName(
                self, 
                "Сохранить отчет в PDF", 
                "", 
                "PDF файлы (*.pdf)"
            )
            
            if not filename:
                return

            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            registerFontFamily('DejaVuSans', normal='DejaVuSans', bold='DejaVuSans')

            doc = SimpleDocTemplate(filename, pagesize=A4)
            elements = []

            header_style = ParagraphStyle(
                'Header',
                fontName='DejaVuSans',
                fontSize=16,
                alignment=TA_CENTER,
                spaceAfter=20
            )

            category_style = ParagraphStyle(
                'Category',
                fontName='DejaVuSans',
                fontSize=14,
                spaceBefore=12,
                spaceAfter=6
            )

            payment_style = ParagraphStyle(
                'Payment',
                fontName='DejaVuSans',
                fontSize=12,
                leading=16
            )

            period_text = f"Список платежей с {report_data['date_from'].strftime('%d.%m.%Y')} по {report_data['date_to'].strftime('%d.%m.%Y')}"
            elements.append(Paragraph(period_text, header_style))

            for category_name, payments_list in report_data['grouped_payments'].items():
                elements.append(Paragraph(f"<b>{category_name}</b>", category_style))
                for p in payments_list:
                    date_str = p.дата.strftime('%d.%m.%Y')
                    elements.append(Paragraph(
                        f"{date_str} — {p.наименование_платежа} ............. {p.стоимость:.2f} р.",
                        payment_style
                    ))
                elements.append(Spacer(1, 10))

            elements.append(Spacer(1, 20))
            elements.append(Paragraph(f"<b>ИТОГО:</b> {report_data['total_sum']:.2f} р.", category_style))

            def add_footer(canvas: Canvas, doc):
                fio = f"{report_data['user'].фамилия} {report_data['user'].имя} {report_data['user'].отчество or ''}"
                canvas.saveState()
                canvas.setFont("DejaVuSans", 10)
                canvas.drawRightString(200 * mm, 10 * mm, f"Страница {doc.page}")
                canvas.drawString(20 * mm, 10 * mm, fio.strip())
                canvas.restoreState()

            doc.build(elements, onFirstPage=add_footer, onLaterPages=add_footer)

            QMessageBox.information(self, "Готово", "PDF отчет успешно сохранен.")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения PDF:\n{str(e)}")
            
    def save_to_docx(self, report_data):
        try:
            filename, _ = QFileDialog.getSaveFileName(
                self, 
                "Сохранить отчет в DOCX", 
                "", 
                "Word документы (*.docx)"
            )
            
            
            if not filename:
                return

            doc = Document()
            
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            title = doc.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run("Отчет о платежах")
            title_run.bold = True
            title_run.font.size = Pt(14)
            
            period = doc.add_paragraph()
            period.alignment = WD_ALIGN_PARAGRAPH.CENTER
            period.add_run(
                f"Период: {report_data['date_from'].strftime('%d.%m.%Y')} - {report_data['date_to'].strftime('%d.%m.%Y')}"
            )
            
            doc.add_paragraph()  
            
            for category, payments in report_data['grouped_payments'].items():
                cat_para = doc.add_paragraph()
                cat_run = cat_para.add_run(category)
                cat_run.bold = True
                
                for payment in payments:
                    para = doc.add_paragraph()
                    para.add_run(f"{payment.дата.strftime('%d.%m.%Y')} - {payment.наименование_платежа}")
                    
                    para.add_run("\t").bold = True
                    last_run = para.add_run(f"{payment.стоимость:.2f} р.")
                    last_run.bold = True
                
                doc.add_paragraph()  
            
     
            total_para = doc.add_paragraph()
            total_run = total_para.add_run(f"ИТОГО: {report_data['total_sum']:.2f} р.")
            total_run.bold = True
            
            doc.add_paragraph("\n")  
            sign_para = doc.add_paragraph()
            sign_para.add_run(
                f"{report_data['user'].фамилия} {report_data['user'].имя} {report_data['user'].отчество or ''}"
            )
            
            doc.save(filename)
            QMessageBox.information(self, "Готово", "DOCX отчет успешно сохранен.")

        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка сохранения DOCX:\n{str(e)}")





        
        

if __name__ == "__main__":
    app = QApplication(sys.argv)
        
    login_window = LoginWindow()
    if login_window.exec() == QDialog.DialogCode.Accepted:
        main_window = PaymentApp(login_window.current_user)
        main_window.show()
        sys.exit(app.exec())
    else:
      sys.exit()
        